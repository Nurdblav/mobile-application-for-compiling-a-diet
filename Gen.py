import random
from unicodedata import name
import copy
import pandas as pd
import Dishes
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.dml import MSO_THEME_COLOR_INDEX
import time

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink

# Класс в котором содержится информация о пользователе
class User:
    sex = ""
    age = 0
    weight = 0
    height = 0
    calorieNorm = 0
    phys_activiti_coef = 0
    

    def __init__(self, sex, age, weight, height, coef):
        self.sex = sex
        self.age = age
        self.weight = weight
        self.height = height
        self.phys_activiti_coef = coef
        if (sex == 'Муж'):
            self.calorieNorm = ((10*self.weight) + (6.25*self.height) - (5*self.age) + 5)*self.phys_activiti_coef
        else:
           self.calorieNorm = ((10*self.weight) + (6.25*self.height) - (5*self.age) - 161)*self.phys_activiti_coef 


# Класс, в которов представлен дневной рацион
class DaysCalorage:
    breakfast = []
    lunch = []
    dinner = []
    totalCalorage = 0

    def __init__ (self, breakfast, lunch, dinner):
        self.breakfast = breakfast
        self.lunch = lunch
        self.dinner = dinner

    def calculateTotal (self):
        self.totalCalorage = 0
        for i in self.breakfast:
            self.totalCalorage += i.calorage_per_100 * i.quantity
        for i in self.lunch:
            self.totalCalorage += i.calorage_per_100 * i.quantity
        for i in self.dinner:
            self.totalCalorage += i.calorage_per_100 * i.quantity


# Класс, который представляет рацион питания на неделю (вариант решения)
class Solution:
    days = []
    totalCalorage = 0
    fitnessFun = 0

    def __init__ (self, days):
        self.days = days
        for i in days:
            self.totalCalorage += i.totalCalorage

    def calculateTotal (self):
        self.totalCalorage = 0
        for i in self.days:
            self.totalCalorage += i.totalCalorage


# Класс в котором реализована работа генетического алгоритма
class GeneticAlgorithm:
    population = []
    populationSize = 0
    iterationsQuantity = 0
    mutationProbability = 0
    bestSolution = Solution # изменить определение
    user = User
    dishes = Dishes.DishesList

    def __init__ (self, populationSize, iterationsQuantity, mutationProbability, user):
        self.populationSize = populationSize
        self.iterationsQuantity = iterationsQuantity
        self.mutationProbability = mutationProbability
        self.user = user
        self.dishes = Dishes.DishesList('dishes.txt')


    def getBreakfast (self):
        tempBreakfast = []
        tempTotalCalorage = 0
        while self.user.calorieNorm*0.35 > tempTotalCalorage:
            tempDish = self.dishes.getDish()
            tempBreakfast.append(tempDish)
            tempTotalCalorage += tempDish.calorage_per_100*tempDish.quantity
        return tempBreakfast

    def getLunch (self):
        tempLunch = []
        tempTotalCalorage = 0
        while self.user.calorieNorm*0.4 > tempTotalCalorage:
            tempDish = self.dishes.getDish()
            tempLunch.append(tempDish)
            tempTotalCalorage += tempDish.calorage_per_100*tempDish.quantity
        return tempLunch

    def getDinner (self):
        tempDinner = []
        tempTotalCalorage = 0
        while self.user.calorieNorm*0.25 > tempTotalCalorage:
            tempDish = self.dishes.getDish()
            tempDinner.append(tempDish)
            tempTotalCalorage += tempDish.calorage_per_100*tempDish.quantity
        return tempDinner

    def initialization (self):
        for i in range(0, self.populationSize):
            tempDaysList = []
            for j in range (0,7):
                tempDay = DaysCalorage(self.getBreakfast(), self.getLunch(), self.getDinner())
                tempDay.calculateTotal()
                tempDaysList.append(tempDay)
            self.population.append(Solution(tempDaysList))
        self.fitnessFun()
        self.bestSolution = copy.deepcopy(self.population[0])


    def fitnessFun (self):
        for i in self.population:
            i.fitnessFun = abs(self.user.calorieNorm*7 - i.totalCalorage)
        self.population.sort(key = lambda x: x.fitnessFun)
        if (self.bestSolution.fitnessFun > self.population[0].fitnessFun):
            self.bestSolution = copy.deepcopy(self.population[0])

    def tournament (self):
        first_part = []
        second_part = [] 
        for i in range(0,5):
            first_part.append(self.population[random.randint(0,len(self.population)-1)])
            second_part.append(self.population[random.randint(0,len(self.population)-1)])
        first_part.sort(key = lambda x: x.fitnessFun)
        second_part.sort(key = lambda x: x.fitnessFun)
        return copy.deepcopy(first_part[0]), copy.deepcopy(second_part[0])

    def crossover (self):
        newPopulation = []
        firstParent = Solution([])
        secondParent = Solution([])
        while len(self.population) > len(newPopulation):
            firstPoint = random.randint(0,3)
            secondPoint = random.randint(4,6)
            firstParent, secondParent = self.tournament()
            child1 = copy.deepcopy(firstParent)
            child2 = copy.deepcopy(secondParent)
            for i in range(firstPoint, secondPoint+1):
                child1.days[i] = copy.deepcopy(secondParent.days[i])
                child2.days[i] = copy.deepcopy(firstParent.days[i])
            child1.calculateTotal()
            child2.calculateTotal()
            newPopulation.append(child1)
            newPopulation.append(child2)
        while len(newPopulation) > self.populationSize:
            newPopulation.pop()
        self.population = copy.deepcopy(newPopulation)

    def mutation(self):
        for i in range(0, len(self.population)):
            if (random.randint(1,100) < self.mutationProbability):
                r_day = random.randint(0, 6)
                if (i < len(self.population)/2):
                    r_intake = random.randint(0, 2)
                    if (r_intake == 0):
                        tempBr = self.getBreakfast()
                        self.population[i].days[r_day].breakfast = tempBr
                        self.population[i].days[r_day].calculateTotal()
                    if (r_intake == 1):
                        tempLun = self.getLunch()
                        self.population[i].days[r_day].lunch = tempLun
                        self.population[i].days[r_day].calculateTotal()
                    if (r_intake == 2):
                        tempDin = self.getDinner()
                        self.population[i].days[r_day].dinner = tempDin
                        self.population[i].days[r_day].calculateTotal()
                else:
                    tempBr = self.getBreakfast()
                    tempLun = self.getLunch()
                    tempDin = self.getDinner()
                    self.population[i].days[r_day].breakfast = tempBr
                    self.population[i].days[r_day].lunch = tempLun
                    self.population[i].days[r_day].dinner = tempDin
                    self.population[i].days[r_day].calculateTotal()
                self.population[i].calculateTotal()

    
    def GeneticAlg (self):
        start_time = time.time()
        self.initialization()
        for i in range(0, self.iterationsQuantity):
            self.crossover()
            self.mutation()
            self.fitnessFun()
        # print("--- %s seconds ---" % (time.time() - start_time))
        print(self.bestSolution.fitnessFun)
        #вывод решения
        # self.printSolution()

    def printSolution (self):
        document = Document()
        style = document.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(14)
        head = document.add_heading('')
        heading = head.add_run('Меню на ближайшие 7 дней.')
        heading.font.size = Pt(20)
        heading.bold = True
        
        for i in range(1,8):
            day_p = document.add_paragraph('День №'+str(i))
            day_p.bold = True
            meal_count = 1
            document.add_paragraph('Завтрак:')
            for j in self.bestSolution.days[i-1].breakfast:
                meal_p = document.add_paragraph(str(meal_count) + '. ')
                add_hyperlink(meal_p, j.name, j.href)
                meal_p.add_run(' ' + str(j.quantity) + '00 грамм - ' + str(round(j.calorage_per_100*j.quantity,2)) + ' калорий')
                meal_count += 1
            document.add_paragraph('Обед:')
            for j in self.bestSolution.days[i-1].lunch:
                meal_p = document.add_paragraph(str(meal_count) + '. ')
                add_hyperlink(meal_p, j.name, j.href)
                meal_p.add_run(' ' + str(j.quantity) + '00 грамм - ' + str(round(j.calorage_per_100*j.quantity,2)) + ' калорий')
                meal_count += 1
            document.add_paragraph('Ужин:')
            for j in self.bestSolution.days[i-1].dinner:
                meal_p = document.add_paragraph(str(meal_count) + '. ')
                add_hyperlink(meal_p, j.name, j.href)
                meal_p.add_run(' ' + str(j.quantity) + '00 грамм - ' + str(round(j.calorage_per_100*j.quantity,2)) + ' калорий')
                meal_count += 1
            document.add_paragraph('Всего калорий за день: ' + str(round(self.bestSolution.days[i-1].totalCalorage,2)))
        
        document.add_paragraph('Всего калорий за 7 дней: ' + str(round(self.bestSolution.totalCalorage,2)))
        document.save('menu.docx')

            
            

